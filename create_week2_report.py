import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, 'F1F5F9')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()

doc = docx.Document()

for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

title_p = doc.add_paragraph()
title_run = title_p.add_run("WEEK 2 TASK REPORT: DEVELOPING INTERACTIVE UI COMPONENTS")
title_run.font.size = Pt(18); title_run.font.bold = True
title_run.font.color.rgb = RGBColor(79, 70, 229)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("Project Name: ComponentHub — Interactive UI Component System | Technologies: HTML5, CSS3, Vanilla JS")
sub_run.font.size = Pt(10); sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 116, 139)
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_heading("1. Executive Summary & Component Requirements", level=1)
doc.add_paragraph(
    "The focus of Week 2 was to design, develop, and test modular, interactive front-end components using pure Vanilla JavaScript, "
    "semantic HTML5, and modern CSS3. ComponentHub comprises four distinct interactive UI systems: "
    "1) Accessible Modal Dialog System, 2) Animated Accordion System, 3) Dynamic Tabbed Panel System, and 4) Toast Notification Manager. "
    "All components strictly enforce ARIA accessibility guidelines, focus management, keyboard arrow navigation, and responsive CSS Grid/Flexbox design."
)

doc.add_heading("2. Concrete Technical Implementation & Code Evidence", level=1)
doc.add_paragraph("Below are exact source code snippets demonstrating DOM manipulation, event handling, animations, and ARIA state updates:")

# Snippet 1: Modal Focus Trap
doc.add_heading("Snippet 2.1: Accessible Modal Dialog & Focus Trap (script.js)", level=2)
code_modal = (
    'function openModal(titleText, bodyText) {\n'
    '    previouslyFocusedElement = document.activeElement;\n'
    '    modalBackdrop.classList.add("active");\n'
    '    modalBackdrop.setAttribute("aria-hidden", "false");\n'
    '    setTimeout(() => closeModalBtn.focus(), 100);\n'
    '    document.addEventListener("keydown", trapModalFocus);\n'
    '}\n\n'
    'function trapModalFocus(e) {\n'
    '    if (e.key === "Escape") { closeModal(); return; }\n'
    '    if (e.key === "Tab") {\n'
    '        const focusables = modalBackdrop.querySelectorAll("button, [href], input, [tabindex]:not([tabindex=\\"-1\\"])");\n'
    '        const first = focusables[0]; const last = focusables[focusables.length - 1];\n'
    '        if (e.shiftKey && document.activeElement === first) { last.focus(); e.preventDefault(); }\n'
    '        else if (!e.shiftKey && document.activeElement === last) { first.focus(); e.preventDefault(); }\n'
    '    }\n'
    '}'
)
create_code_block(doc, code_modal)
doc.add_paragraph("Technical Explanation: Focus trapping restricts keyboard navigation within the dialog container while open. The previously active element receives focus restoration upon modal close.")

# Snippet 2: Accordion
doc.add_heading("Snippet 2.2: Animated Accordion & Arrow Key Switching (script.js)", level=2)
code_acc = (
    'header.addEventListener("keydown", (e) => {\n'
    '    if (e.key === "ArrowDown") {\n'
    '        e.preventDefault();\n'
    '        const nextHeader = accordionHeaders[(index + 1) % accordionHeaders.length];\n'
    '        nextHeader.focus();\n'
    '    } else if (e.key === "ArrowUp") {\n'
    '        e.preventDefault();\n'
    '        const prevHeader = accordionHeaders[(index - 1 + accordionHeaders.length) % accordionHeaders.length];\n'
    '        prevHeader.focus();\n'
    '    }\n'
    '});'
)
create_code_block(doc, code_acc)
doc.add_paragraph("Technical Explanation: Listens for ArrowUp and ArrowDown key presses to cycle active focus between accordion headers seamlessly.")

# Snippet 3: Tabs
doc.add_heading("Snippet 2.3: Dynamic Tabbed Panels & W3C ARIA Roles (index.html & script.js)", level=2)
code_tabs = (
    '<div class="tab-list" role="tablist" aria-label="Component Information Tabs">\n'
    '    <button role="tab" aria-selected="true" aria-controls="panel-overview" id="tab-overview" tabindex="0" class="tab-btn active">📊 Overview</button>\n'
    '    <button role="tab" aria-selected="false" aria-controls="panel-analytics" id="tab-analytics" tabindex="-1" class="tab-btn">📈 Analytics</button>\n'
    '</div>'
)
create_code_block(doc, code_tabs)
doc.add_paragraph("Technical Explanation: Implements W3C WAI-ARIA authoring practices for tabs, updating tabindex (0 / -1) and aria-selected states during tab selection.")

doc.add_heading("3. Component State Verification & Testing Matrix", level=1)
doc.add_paragraph("Comprehensive state testing results across all four UI components:")

table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
headers = ["Component Name", "Interactive States Verified", "ARIA Attributes Synchronized", "Testing Result"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_background(hdr_cells[i], '4F46E5')
    p = hdr_cells[i].paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255); p.runs[0].font.bold = True

test_data = [
    ("Modal Dialog System", "Open animation, Focus trap, ESC key dismiss, Backdrop click close", "role='dialog', aria-modal='true', aria-hidden", "PASSED ✅"),
    ("Animated Accordion", "Expand/Collapse, CSS max-height transition, Arrow key focus", "aria-expanded, aria-controls, aria-labelledby", "PASSED ✅"),
    ("Tabbed Panels", "Panel swapping, Sliding indicator, Left/Right arrow switching", "role='tablist', role='tab', role='tabpanel', aria-selected", "PASSED ✅"),
    ("Toast Notifications", "Dynamic DOM spawning, Auto-dismiss timer (4s), Manual dismiss", "role='status', aria-live='polite', aria-atomic", "PASSED ✅")
]

for row_idx, data in enumerate(test_data, start=1):
    row_cells = table.rows[row_idx].cells
    for col_idx, text in enumerate(data):
        row_cells[col_idx].text = text
        if col_idx == 3: set_cell_background(row_cells[col_idx], 'D1FAE5')

doc.add_paragraph()

doc.add_heading("4. Deliverable Files & Package Verification", level=1)
doc.add_paragraph(
    "All component files (index.html, style.css, script.js, README.md) and the deliverable compressed archive "
    "(interactive_ui_components.zip) are packaged in the repository root and pushed to GitHub: "
    "https://github.com/Zainul9142/YuvaIntern2"
)

doc.save("Interactive_UI_Components_Week2_Report.docx")
print("Interactive_UI_Components_Week2_Report.docx generated successfully!")
